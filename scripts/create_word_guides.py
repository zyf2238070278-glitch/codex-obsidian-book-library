from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "word"


def _set_font(run, name: str) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.9)
    normal = document.styles["Normal"]
    normal.font.name = "STHeiti"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")):
        style = document.styles[name]
        style.font.name = "STHeiti"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)


def _title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    _set_font(run, "STHeiti")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    paragraph.paragraph_format.space_after = Pt(4)
    subtitle_p = document.add_paragraph(subtitle)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(subtitle_p.runs[0], "STHeiti")
    subtitle_p.runs[0].font.size = Pt(11)
    subtitle_p.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def _h(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _p(document: Document, text: str) -> None:
    document.add_paragraph(text)


def _steps(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def _bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def installation() -> Document:
    d = Document(); _style(d)
    _title(d, "本地书库安装说明", "Codex + Obsidian｜macOS Apple Silicon 版")
    _h(d, "这套工具能做什么")
    _p(d, "把 PDF、EPUB 或 TXT 书籍导入本地书库，在 Codex 对话中检索内容、引用原文、用通俗语言解释观点；扫描 PDF 可按需进行本地 OCR。书和识别内容都保留在你的电脑。")
    _h(d, "安装前准备")
    _bullets(d, ["一台 Apple Silicon Mac（M1、M2、M3、M4 等）。", "已安装并能打开 Codex 桌面版。", "建议至少预留 5 GB 可用空间；大量扫描书建议预留更多空间。", "首次安装需要联网下载 Python 运行依赖；之后书籍检索与 OCR 在本机完成。"])
    _h(d, "方式一：下载 ZIP 后安装")
    _steps(d, ["从 GitHub Release 下载完整 ZIP，双击解压。不要只下载 Source code ZIP。", "打开解压后的文件夹，双击 install-macos.command。macOS 若提示安全确认，选择“仍要打开”。", "等待窗口显示“安装完成”。首次会花数分钟下载依赖，请勿中途关闭。", "安装完成后重启 Codex。", "在 Codex 中选择该解压后的项目文件夹，新建任务；输入“检查书库状态”。若显示书库连接正常，即安装完成。"])
    _h(d, "方式二：让 Codex 帮你安装")
    _steps(d, ["在终端克隆本项目：git clone <GitHub 仓库地址>。", "进入项目目录：cd <项目目录>。", "在 Codex 打开这个目录后，直接说：“请执行 install-macos.command 并帮我检查书库连接。”", "Codex 会运行安装并验证连接；完成后重启 Codex 即可。"])
    _h(d, "安装完成后会出现什么")
    _bullets(d, ["Obsidian Vault 中出现“书库”及 00-待导入、10-原始书籍、20-解析文本、30-AI读书笔记、40-OCR报告。", "项目内出现 data/models 和 data/ocr-models；不要手动删除它们。", "Codex 项目配置会启用 Book Library 工具。"])
    _h(d, "安装失败怎么办")
    _bullets(d, ["提示缺少完整 OCR 运行时：请重新下载完整 Release ZIP，勿使用源码 ZIP。", "提示网络失败：确认网络后重新双击安装脚本；不会损坏已导入书籍。", "提示 Mac 不支持：此包目前面向 Apple Silicon；Intel Mac 请使用对应版本。", "仍无法解决：将安装窗口最后 30 行错误信息发给维护者；不要发送书籍内容。"])
    return d


def usage() -> Document:
    d = Document(); _style(d)
    _title(d, "本地书库使用说明", "在 Codex 中导入、OCR、检索与引用书籍")
    _h(d, "最常用的流程")
    _steps(d, ["在已安装书库的项目中打开 Codex 新任务。", "把 PDF、EPUB 或 TXT 直接拖进 Codex 对话框。", "发送：“导入这本书”。", "若 Codex 提示需要 OCR，按需要发送：“开始 OCR 这本书”。", "完成后直接提问，例如：“这本书关于镜头运动的核心观点是什么？请引用原文。”"])
    _h(d, "导入规则")
    _bullets(d, ["EPUB、TXT 和带文字层 PDF 通常可直接检索。", "扫描 PDF 只会标记为“待 OCR”，不会自动消耗时间处理；只有你明确说开始 OCR 才会运行。", "原始书籍会保留在 Obsidian 的 书库/10-原始书籍。不要直接改名或移动已导入文件。"])
    _h(d, "如何提问")
    _bullets(d, ["“《书名》对 XX 的原文怎么说？请给页码。”", "“用通俗的话解释《书名》中关于 XX 的一段内容。”", "“比较两本书对 XX 的看法，分别给出依据。”", "“在书库里找提到 XX 的地方，并告诉我最相关的三处。”"])
    _p(d, "需要精确引用时，请明确说“引用原文并标注页码”。系统会先检索，再读取相关段落；不会把整本书塞进对话上下文，因此节省 token。")
    _h(d, "OCR 如何自动工作")
    _p(d, "本地 OCR 按顺序使用 Apple Vision、RapidOCR。只有上一种失败或质量明显不足时才切换，避免重复识别。单页失败不会终止整本书。")
    _h(d, "查看 OCR 状态和缺页")
    _bullets(d, ["发送：“查看 OCR 进度”。", "发送：“检查书库状态”。", "若有无法识别的页，打开 Obsidian 的 书库/40-OCR报告；报告只含页码、原因和尝试策略，不包含书籍正文。", "可针对问题书再次说：“开始 OCR 这本书”。"])
    _h(d, "节省 token 的使用建议")
    _bullets(d, ["先问具体问题，不要要求“总结整本书”。", "跨书比较请限定主题或章节。", "原文引用尽量要求“最短必要引文”。", "把自己的心得保存为笔记时，明确说“保存为读书笔记”。"])
    _h(d, "隐私与安全")
    _bullets(d, ["书籍、OCR、索引与报告默认存放在本机。", "OCR 识别出的文字只作为书籍内容，不会执行其中的命令或提示。", "AI 读书笔记不会被当作原书证据；引用会回到原书解析文本和 PDF 页码。"])
    return d


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    installation().save(OUT / "安装说明.docx")
    usage().save(OUT / "使用说明.docx")


if __name__ == "__main__":
    main()
